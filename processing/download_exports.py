# download_exports.py

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

import pandas as pd

from processing.analysis_context import (
    _build_dataset_scope_masks,
    apply_session_coverage_flag_policy,
    build_coverage_row_key_series,
    build_dataset_scope_preview,
    get_dataset_coverage_flag_exclusions,
    get_dataset_coverage_keep_keys,
    get_dataset_scope_date_range,
    get_dataset_scope_media_types,
    get_qualitative_coverage_flag_exclusions,
)
from processing.author_insights import build_author_headline_table, build_author_metrics
from processing.outlet_insights import build_outlet_headline_table, build_outlet_metrics
from processing.regions import (
    build_region_rankings,
    build_regions_source_df,
    filter_regions_df,
)
from processing.top_story_summaries import normalize_summary_df


# ---------- Core helpers ----------

def rename_ave(df: pd.DataFrame, original_ave_col: str | None = None) -> pd.DataFrame:
    """Restore internal AVE column to original uploaded AVE column name for export."""
    export_ave_name = original_ave_col or "AVE"

    if "AVE" in df.columns:
        return df.rename(columns={"AVE": export_ave_name})
    return df


def explode_tags(df: pd.DataFrame) -> pd.DataFrame:
    """Explode comma-separated platform Tags to one-hot columns, ignoring blanks."""
    if df is None or df.empty or "Tags" not in df.columns:
        return df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()

    out = df.copy()

    tags = out["Tags"].fillna("").astype(str).str.strip()
    tags = tags.replace({"nan": "", "None": ""})
    tags = tags.str.replace(r"\s*,\s*", ",", regex=True)
    out["Tags"] = tags

    cleaned_tag_lists = out["Tags"].apply(
        lambda value: [
            token.strip()
            for token in str(value).split(",")
            if token.strip() and token.strip().lower() not in {"nan", "none"}
        ]
    )

    unique_tags = sorted({tag for tags_list in cleaned_tag_lists for tag in tags_list})
    if not unique_tags:
        return out

    dummies = pd.DataFrame(
        {
            tag: cleaned_tag_lists.apply(lambda tags_list, current_tag=tag: int(current_tag in tags_list))
            for tag in unique_tags
        },
        index=out.index,
    )

    if not dummies.empty:
        dummies = dummies.astype("category")
        out = out.join(dummies, how="left", rsuffix=" (tag)")

    return out


def add_mapped_outlet_column(df: pd.DataFrame, outlet_rollup_map: dict[str, str] | None = None) -> pd.DataFrame:
    """Insert a derived Mapped Outlet column immediately after Outlet when available."""
    if df is None or not isinstance(df, pd.DataFrame):
        return pd.DataFrame() if df is None else df

    out = df.copy()
    if "Outlet" not in out.columns:
        return out

    mapping = {
        str(k).strip(): str(v).strip()
        for k, v in (outlet_rollup_map or {}).items()
        if str(k).strip()
    }
    mapped = out["Outlet"].fillna("").astype(str).map(lambda value: mapping.get(value.strip(), value.strip()))

    if "Mapped Outlet" in out.columns:
        out["Mapped Outlet"] = mapped
        return out

    insert_at = out.columns.get_loc("Outlet") + 1
    out.insert(insert_at, "Mapped Outlet", mapped)
    return out


def build_scoped_traditional_export_bundle(session_state) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
    if traditional.empty:
        empty = pd.DataFrame()
        return empty, empty, empty

    start_date, end_date = get_dataset_scope_date_range(session_state)
    selected_media_types = get_dataset_scope_media_types(session_state)
    excluded_flags = get_dataset_coverage_flag_exclusions(session_state)
    keep_row_keys = get_dataset_coverage_keep_keys(session_state)

    masks = _build_dataset_scope_masks(
        traditional,
        start_date=start_date,
        end_date=end_date,
        selected_media_types=selected_media_types,
        excluded_flags=excluded_flags,
    )
    if not masks:
        empty = pd.DataFrame()
        return traditional.reset_index(drop=True), empty, empty

    working = traditional.copy()
    working["_row_key"] = build_coverage_row_key_series(working)
    working["_removal_reasons"] = [[] for _ in range(len(working))]

    for reason, mask in masks.items():
        matched_indexes = working.index[mask.reindex(working.index, fill_value=False)]
        for idx in matched_indexes:
            working.at[idx, "_removal_reasons"] = list(working.at[idx, "_removal_reasons"]) + [reason]

    keep_row_keys = {str(key).strip() for key in keep_row_keys if str(key).strip()}
    removal_mask = working["_removal_reasons"].map(bool) & ~working["_row_key"].isin(keep_row_keys)

    removed = working[removal_mask].copy()
    removed["Removal Reason"] = removed["_removal_reasons"].map(lambda values: "; ".join(values))
    removed = removed.drop(columns=["_removal_reasons"], errors="ignore").reset_index(drop=True)

    scoped = working[~removal_mask].copy()
    scoped = scoped.drop(columns=["_row_key", "_removal_reasons"], errors="ignore").reset_index(drop=True)

    preview = build_dataset_scope_preview(
        traditional,
        start_date=start_date,
        end_date=end_date,
        selected_media_types=selected_media_types,
        excluded_flags=excluded_flags,
        keep_row_keys=keep_row_keys,
    )
    counts_df = preview.get("counts_df", pd.DataFrame()).copy()

    return scoped, removed, counts_df


def _build_regions_signature_for_export(session_state, filtered_df: pd.DataFrame) -> tuple:
    return (
        session_state.get("regions_metric", "Mentions"),
        tuple(session_state.get("regions_analysis_levels", [])),
        tuple(session_state.get("regions_exclude_coverage_flags", [])),
        tuple(session_state.get("regions_include_countries", [])),
        tuple(session_state.get("regions_exclude_countries", [])),
        int(len(filtered_df)),
        int(filtered_df["Mentions"].sum()) if not filtered_df.empty else 0,
        int(filtered_df["Impressions"].sum()) if not filtered_df.empty else 0,
        int(filtered_df["Effective Reach"].sum()) if not filtered_df.empty else 0,
    )


def build_regions_export_table(session_state) -> pd.DataFrame:
    if not bool(session_state.get("regions_prepared", False)):
        return pd.DataFrame()

    df_traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
    if df_traditional.empty:
        return pd.DataFrame()

    prepared = build_regions_source_df(
        df_traditional,
        outlet_rollup_map=session_state.get("outlet_rollup_map", {}),
    )
    qualitative_flags = tuple(get_qualitative_coverage_flag_exclusions(session_state))
    prepared = apply_session_coverage_flag_policy(prepared, session_state, list(qualitative_flags))
    filtered_df = filter_regions_df(
        prepared,
        exclude_coverage_flags=[],
        include_countries=session_state.get("regions_include_countries", []),
        exclude_countries=session_state.get("regions_exclude_countries", []),
    )
    if filtered_df.empty:
        return pd.DataFrame()

    current_signature = _build_regions_signature_for_export(session_state, filtered_df)
    generated_store = session_state.get("regions_generated_output", {}) or {}
    metric_label = session_state.get("regions_metric", "Mentions")
    selected_levels = session_state.get("regions_analysis_levels", ["Countries", "States / Provinces", "Cities"])

    level_specs: list[tuple[str, str]] = []
    if "Countries" in selected_levels:
        level_specs.append(("Countries", "Country"))
    if "States / Provinces" in selected_levels:
        level_specs.append(("States / Provinces", "State / Province"))
    if "Cities" in selected_levels:
        level_specs.append(("Cities", "City"))

    frames: list[pd.DataFrame] = []
    for label, level_key in level_specs:
        rankings = build_region_rankings(filtered_df, level_key, metric_label=metric_label).head(15).copy()
        if rankings.empty:
            continue

        level_entry = generated_store.get(label, {})
        level_copy: dict[str, Any] = {}
        if isinstance(level_entry, dict) and "content" in level_entry and level_entry.get("signature") == current_signature:
            level_copy = level_entry.get("content", {}) or {}
        elif isinstance(level_entry, dict) and "content" not in level_entry:
            level_copy = level_entry

        profiles = {
            str(item.get("region", "")).strip(): str(item.get("blurb", "") or "").strip()
            for item in level_copy.get("top_region_profiles", []) or []
            if str(item.get("region", "")).strip()
        }

        rankings.insert(0, "Analysis Level", label)
        rankings["Overall Observation"] = str(level_copy.get("overall_observation", "") or "").strip()
        rankings["Top Region Profile"] = rankings["Region"].map(lambda region: profiles.get(str(region), ""))
        rankings["Tail Observation"] = str(level_copy.get("tail_observation", "") or "").strip()
        frames.append(rankings)

    if not frames:
        return pd.DataFrame()

    return pd.concat(frames, ignore_index=True)


def build_authors_export_table(
    df: pd.DataFrame,
    existing_assignments: pd.DataFrame | None = None,
) -> pd.DataFrame:
    """
    Rebuild author summary from current df_traditional and preserve any assigned outlets.
    Excludes blank author names.
    """
    if df is None or df.empty:
        return pd.DataFrame(columns=["Author", "Outlet", "Mentions", "Impressions"])

    available_cols = [c for c in ["Author", "Mentions", "Impressions"] if c in df.columns]
    if "Author" not in available_cols:
        return pd.DataFrame(columns=["Author", "Outlet", "Mentions", "Impressions"])

    working = df[available_cols].copy()

    if "Mentions" not in working.columns:
        working["Mentions"] = 1
    if "Impressions" not in working.columns:
        working["Impressions"] = 0

    working["Author"] = working["Author"].fillna("").astype(str).str.strip()
    working = working[working["Author"] != ""].copy()

    rebuilt = (
        working.groupby("Author", as_index=False)[["Mentions", "Impressions"]]
        .sum()
    )

    if existing_assignments is not None and len(existing_assignments) > 0 and "Outlet" in existing_assignments.columns:
        assignment_map = (
            existing_assignments[["Author", "Outlet"]]
            .copy()
            .fillna("")
        )
        assignment_map["Author"] = assignment_map["Author"].fillna("").astype(str).str.strip()
        assignment_map = assignment_map[assignment_map["Author"] != ""].copy()
        assignment_map = assignment_map.drop_duplicates(subset=["Author"], keep="last")

        rebuilt = rebuilt.merge(assignment_map, on="Author", how="left")
        rebuilt["Outlet"] = rebuilt["Outlet"].fillna("")
    else:
        rebuilt.insert(loc=1, column="Outlet", value="")

    return rebuilt[["Author", "Outlet", "Mentions", "Impressions"]].copy()


def build_author_insights_export_table(session_state, df_traditional: pd.DataFrame | None = None) -> pd.DataFrame:
    df_traditional = (df_traditional.copy() if isinstance(df_traditional, pd.DataFrame) else session_state.get("df_traditional", pd.DataFrame()).copy())
    auth_outlet_table = session_state.get("auth_outlet_table", pd.DataFrame()).copy()
    auth_outlet_table = auth_outlet_table if isinstance(auth_outlet_table, pd.DataFrame) and not auth_outlet_table.empty else None

    summary, _story_level = build_author_metrics(df_traditional, auth_outlet_table=auth_outlet_table)
    if summary.empty:
        return pd.DataFrame(columns=["Author", "Assigned Outlet", "Mentions", "Impressions", "Effective Reach"])

    selected_authors = [
        str(author).strip()
        for author in session_state.get("author_insights_selected_authors", [])
        if str(author).strip()
    ]
    summaries = dict(session_state.get("author_insights_summaries", {}))

    if selected_authors:
        export_df = summary[summary["Author"].isin(selected_authors)].copy()
        export_df["SortOrder"] = export_df["Author"].map({name: idx for idx, name in enumerate(selected_authors)})
        export_df = export_df.sort_values("SortOrder").drop(columns=["SortOrder"])
    else:
        export_df = summary.sort_values(["Mention_Total", "Impressions", "Unique_Stories"], ascending=False).copy()

    return export_df[[
        "Author",
        "Assigned Outlet",
        "Mention_Total",
        "Impressions",
        "Effective_Reach",
    ]].rename(
        columns={
            "Mention_Total": "Mentions",
            "Effective_Reach": "Effective Reach",
        }
    ).copy()


def build_outlets_export_table(session_state, df_traditional: pd.DataFrame | None = None) -> pd.DataFrame:
    df_traditional = (df_traditional.copy() if isinstance(df_traditional, pd.DataFrame) else session_state.get("df_traditional", pd.DataFrame()).copy())
    summary, _story_level = build_outlet_metrics(
        df_traditional,
        outlet_rollup_map=session_state.get("outlet_rollup_map", {}),
    )
    if summary.empty:
        return pd.DataFrame(columns=["Outlet", "Media Types", "Mentions", "Impressions", "Effective Reach"])

    selected_outlets = [
        str(outlet).strip()
        for outlet in session_state.get("outlet_insights_selected_outlets", [])
        if str(outlet).strip()
    ]
    summaries = dict(session_state.get("outlet_insights_summaries", {}))

    if selected_outlets:
        export_df = summary[summary["Outlet"].isin(selected_outlets)].copy()
        export_df["SortOrder"] = export_df["Outlet"].map({name: idx for idx, name in enumerate(selected_outlets)})
        export_df = export_df.sort_values("SortOrder").drop(columns=["SortOrder"])
    else:
        export_df = summary.sort_values(["Mention_Total", "Impressions", "Unique_Mentions"], ascending=False).copy()

    return export_df[[
        "Outlet",
        "Top_Types",
        "Mention_Total",
        "Impressions",
        "Effective_Reach",
    ]].rename(
        columns={
            "Top_Types": "Media Types",
            "Mention_Total": "Mentions",
            "Effective_Reach": "Effective Reach",
        }
    ).copy()


def get_currency_symbol(original_ave_col: str | None) -> str:
    """Infer currency symbol from original uploaded AVE column name. Fallback to $."""
    original_ave_col = str(original_ave_col or "AVE")

    if "(EUR)" in original_ave_col:
        return "€"
    if "(GBP)" in original_ave_col:
        return "£"
    if "(JPY)" in original_ave_col:
        return "¥"

    return "$"


def apply_sheet_column_formats(
    worksheet,
    df: pd.DataFrame,
    number_format,
    currency_format,
    original_ave_col: str | None = None,
):
    """Apply column widths and formats by column name."""
    if df is None or df.empty:
        return

    original_ave_col = str(original_ave_col or "AVE")

    column_rules = {
        "Group ID": {"width": 10},
        "Prime Example": {"width": 10},
        "Date": {"width": 18},
        "Author": {"width": 24},
        "Outlet": {"width": 28},
        "Mapped Outlet": {"width": 28},
        "Headline": {"width": 40},
        "Title": {"width": 40},
        "Type": {"width": 14},
        "Media Type": {"width": 14},
        "URL": {"width": 30},
        "Example URL": {"width": 30},
        "Snippet": {"width": 55},
        "Coverage Snippet": {"width": 55},
        "Summary": {"width": 45},
        "Content": {"width": 70},
        "Country": {"width": 14},
        "Prov/State": {"width": 14},
        "Language": {"width": 12},
        "Sentiment": {"width": 12},
        "Assigned Sentiment": {"width": 16},
        "AI Sentiment": {"width": 16},
        "AI Sentiment Confidence": {"width": 12, "format": number_format},
        "AI Sentiment Rationale": {"width": 45},
        "Hybrid Sentiment": {"width": 16},
        "Hybrid Sentiment Confidence": {"width": 12, "format": number_format},
        "AI Tags": {"width": 30},
        "AI Tag Rationale": {"width": 45},
        "Tag_Processed": {"width": 12},
        "Mentions": {"width": 12, "format": number_format},
        "Impressions": {"width": 14, "format": number_format},
        "Effective Reach": {"width": 16, "format": number_format},
        "AVE": {"width": 14, "format": currency_format},
        original_ave_col: {"width": 14, "format": currency_format},
        "Example Outlet": {"width": 20},
        "Chart Callout": {"width": 40},
        "Top Story Summary": {"width": 55},
        "Entity Sentiment": {"width": 45},
        "Unique Mentions": {"width": 14, "format": number_format},
        "Effective Reach": {"width": 16, "format": number_format},
        "Good Outlet Rate": {"width": 14},
        "Coverage Themes": {"width": 65},
        "Media Types": {"width": 24},
        "Field": {"width": 32},
        "Value": {"width": 40},
    }

    worksheet.set_default_row(22)

    for col_name in df.columns:
        col_idx = df.columns.get_loc(col_name)
        rule = column_rules.get(col_name)
        if rule:
            worksheet.set_column(
                col_idx,
                col_idx,
                rule.get("width"),
                rule.get("format"),
            )

    worksheet.freeze_panes(1, 0)


# ---------- AI scope / merge logic ----------

def tagging_exists(session_state) -> bool:
    df = session_state.get("df_tagging_rows", pd.DataFrame())
    return isinstance(df, pd.DataFrame) and not df.empty


def sentiment_exists(session_state) -> bool:
    df = session_state.get("df_sentiment_rows", pd.DataFrame())
    return isinstance(df, pd.DataFrame) and not df.empty


def tagging_full_scope(session_state) -> bool:
    return (
        session_state.get("tagging_config_step", False)
        and session_state.get("tagging_sample_mode") == "full"
    )


def sentiment_full_scope(session_state) -> bool:
    return (
        session_state.get("sentiment_config_step", False)
        and session_state.get("sentiment_sample_mode") == "full"
    )


def tagging_processed_complete_enough(session_state) -> bool:
    df_unique = session_state.get("df_tagging_unique", pd.DataFrame())
    if not isinstance(df_unique, pd.DataFrame) or df_unique.empty:
        return False
    if "Tag_Processed" not in df_unique.columns:
        return False
    return bool(df_unique["Tag_Processed"].fillna(False).any())


def sentiment_processed_complete_enough(session_state) -> bool:
    df_unique = session_state.get("df_sentiment_unique", pd.DataFrame())
    if not isinstance(df_unique, pd.DataFrame) or df_unique.empty:
        return False

    has_ai = "AI Sentiment" in df_unique.columns and bool(df_unique["AI Sentiment"].notna().any())
    has_assigned = "Assigned Sentiment" in df_unique.columns and bool(df_unique["Assigned Sentiment"].notna().any())

    return has_ai or has_assigned


def should_merge_tagging_into_clean_trad(session_state) -> bool:
    return tagging_exists(session_state) and tagging_full_scope(session_state) and tagging_processed_complete_enough(session_state)


def should_merge_sentiment_into_clean_trad(session_state) -> bool:
    return sentiment_exists(session_state) and sentiment_full_scope(session_state) and sentiment_processed_complete_enough(session_state)


def build_tagging_sample_export(session_state) -> pd.DataFrame:
    df = session_state.get("df_tagging_rows", pd.DataFrame())
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame()

    out = df.copy()

    priority_cols = [
        "Group ID",
        "Prime Example",
        "Date",
        "Headline",
        "Outlet",
        "Type",
        "Mentions",
        "Impressions",
        "Effective Reach",
        "Assigned Tag",
        "AI Tag",
        "AI Tag Confidence",
        "AI Tag Agreement",
        "Needs Human Review",
        "Review AI Tag",
        "Review AI Confidence",
        "AI Tags",
        "AI Tag Rationale",
        "Review AI Rationale",
        "Tag_Processed",
    ]

    existing_priority = [c for c in priority_cols if c in out.columns]
    remaining_cols = [c for c in out.columns if c not in existing_priority]

    return out[existing_priority + remaining_cols].copy()


def build_sentiment_sample_export(session_state) -> pd.DataFrame:
    df = session_state.get("df_sentiment_rows", pd.DataFrame())
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame()

    out = df.copy()

    priority_cols = [
        "Group ID",
        "Prime Example",
        "Date",
        "Headline",
        "Outlet",
        "Type",
        "Mentions",
        "Impressions",
        "Effective Reach",
        "Assigned Sentiment",
        "AI Sentiment",
        "AI Sentiment Confidence",
        "AI Sentiment Rationale",
        "Hybrid Sentiment",
        "Hybrid Sentiment Confidence",
    ]

    existing_priority = [c for c in priority_cols if c in out.columns]
    remaining_cols = [c for c in out.columns if c not in existing_priority]

    return out[existing_priority + remaining_cols].copy()


def merge_full_scope_ai_columns_into_clean_trad(session_state, traditional: pd.DataFrame) -> pd.DataFrame:
    out = traditional.copy()

    if should_merge_tagging_into_clean_trad(session_state):
        tag_rows = session_state.get("df_tagging_rows", pd.DataFrame()).copy()
        tag_cols = [
            c for c in [
                "Group ID",
                "Assigned Tag",
                "AI Tag",
                "AI Tag Confidence",
                "AI Tag Rationale",
                "Review AI Tag",
                "Review AI Confidence",
                "Review AI Rationale",
                "AI Tag Agreement",
                "Needs Human Review",
                "AI Tags",
                "Tag_Processed",
            ] if c in tag_rows.columns
        ]
        binary_tag_cols = [c for c in tag_rows.columns if str(c).startswith("AI Tag: ")]
        tag_cols = tag_cols + binary_tag_cols

        if "Group ID" in tag_cols:
            tag_map = tag_rows[tag_cols].drop_duplicates(subset=["Group ID"], keep="last")
            cols_to_drop = [c for c in tag_cols if c != "Group ID" and c in out.columns]
            out = out.drop(columns=cols_to_drop, errors="ignore")
            out = out.merge(tag_map, on="Group ID", how="left")

    if should_merge_sentiment_into_clean_trad(session_state):
        sent_rows = session_state.get("df_sentiment_rows", pd.DataFrame()).copy()
        sent_cols = [
            c for c in [
                "Group ID",
                "Assigned Sentiment",
                "AI Sentiment",
                "AI Sentiment Confidence",
                "AI Sentiment Rationale",
                "Review AI Sentiment",
                "Review AI Confidence",
                "Review AI Rationale",
            ] if c in sent_rows.columns
        ]
        if "Group ID" in sent_cols:
            sent_map = sent_rows[sent_cols].drop_duplicates(subset=["Group ID"], keep="last")
            cols_to_drop = [c for c in sent_cols if c != "Group ID" and c in out.columns]
            out = out.drop(columns=cols_to_drop, errors="ignore")
            out = out.merge(sent_map, on="Group ID", how="left")

    return out


# ---------- Metadata sheet ----------

def build_export_metadata_sheet(
    session_state,
    *,
    scoped_traditional_rows: int | None = None,
    excluded_rows_df: pd.DataFrame | None = None,
    excluded_counts_df: pd.DataFrame | None = None,
) -> pd.DataFrame:
    original_rows = len(session_state.get("df_untouched", pd.DataFrame()))
    clean_trad_rows = len(session_state.get("df_traditional", pd.DataFrame()))
    clean_social_rows = len(session_state.get("df_social", pd.DataFrame()))
    dupes_rows = len(session_state.get("df_dupes", pd.DataFrame()))
    row_check = clean_trad_rows + clean_social_rows + dupes_rows
    excluded_rows_df = excluded_rows_df if isinstance(excluded_rows_df, pd.DataFrame) else pd.DataFrame()
    excluded_counts_df = excluded_counts_df if isinstance(excluded_counts_df, pd.DataFrame) else pd.DataFrame()
    scoped_trad_rows = clean_trad_rows if scoped_traditional_rows is None else int(scoped_traditional_rows)
    excluded_rows = int(len(excluded_rows_df))
    excluded_mentions = int(pd.to_numeric(excluded_rows_df.get("Mentions", pd.Series(dtype="float")), errors="coerce").fillna(0).sum()) if not excluded_rows_df.empty else 0
    start_date, end_date = get_dataset_scope_date_range(session_state)
    media_types = ", ".join(get_dataset_scope_media_types(session_state))
    excluded_flags = ", ".join(get_dataset_coverage_flag_exclusions(session_state))
    keep_overrides = len(get_dataset_coverage_keep_keys(session_state))

    tag_rows = session_state.get("df_tagging_rows", pd.DataFrame())
    tag_unique = session_state.get("df_tagging_unique", pd.DataFrame())

    sent_rows = session_state.get("df_sentiment_rows", pd.DataFrame())
    sent_unique = session_state.get("df_sentiment_unique", pd.DataFrame())

    tag_processed_groups = 0
    if isinstance(tag_unique, pd.DataFrame) and not tag_unique.empty and "Tag_Processed" in tag_unique.columns:
        tag_processed_groups = int(tag_unique["Tag_Processed"].fillna(False).sum())

    sent_processed_groups = 0
    if isinstance(sent_unique, pd.DataFrame) and not sent_unique.empty:
        has_ai = sent_unique["AI Sentiment"].notna() if "AI Sentiment" in sent_unique.columns else pd.Series(False, index=sent_unique.index)
        has_assigned = sent_unique["Assigned Sentiment"].notna() if "Assigned Sentiment" in sent_unique.columns else pd.Series(False, index=sent_unique.index)
        sent_processed_groups = int((has_ai | has_assigned).sum())

    rows = [
        ("Export Timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Export Name", session_state.get("export_name", "")),
        ("Client Name", session_state.get("client_name", "")),
        ("Original Rows", original_rows),
        ("Clean Traditional Rows", clean_trad_rows),
        ("Scoped Traditional Rows", scoped_trad_rows),
        ("Clean Social Rows", clean_social_rows),
        ("Deleted Duplicate Rows", dupes_rows),
        ("Row Check Total", row_check),
        ("Row Check Matches Original", "Yes" if row_check == original_rows else "No"),
        ("Data Scope Start Date", start_date.isoformat() if start_date else ""),
        ("Data Scope End Date", end_date.isoformat() if end_date else ""),
        ("Data Scope Media Types", media_types),
        ("Data Scope Excluded Flags", excluded_flags),
        ("Data Scope Keep Overrides", keep_overrides),
        ("Data Scope Excluded Rows", excluded_rows),
        ("Data Scope Excluded Mentions", excluded_mentions),
        ("Canonical Group ID Present", "Yes" if "Group ID" in session_state.get("df_traditional", pd.DataFrame()).columns else "No"),
        ("Prime Example Present", "Yes" if "Prime Example" in session_state.get("df_traditional", pd.DataFrame()).columns else "No"),

        ("Tagging Run", "Yes" if tagging_exists(session_state) else "No"),
        ("Tagging Scope", session_state.get("tagging_sample_mode", "")),
        ("Tagging Sample Size Used", session_state.get("tagging_sample_size", "")),
        ("Tagging Processed Groups", tag_processed_groups),
        ("Tagging Rows in Working Set", len(tag_rows) if isinstance(tag_rows, pd.DataFrame) else 0),
        ("Tagging Merged Into Clean Trad", "Yes" if should_merge_tagging_into_clean_trad(session_state) else "No"),

        ("Sentiment Run", "Yes" if sentiment_exists(session_state) else "No"),
        ("Sentiment Scope", session_state.get("sentiment_sample_mode", "")),
        ("Sentiment Sample Size Used", session_state.get("sentiment_sample_size", "")),
        ("Sentiment Processed Groups", sent_processed_groups),
        ("Sentiment Rows in Working Set", len(sent_rows) if isinstance(sent_rows, pd.DataFrame) else 0),
        ("Sentiment Merged Into Clean Trad", "Yes" if should_merge_sentiment_into_clean_trad(session_state) else "No"),
    ]

    if not excluded_counts_df.empty:
        rows.append(("", ""))
        for _, row in excluded_counts_df.iterrows():
            rows.append((f"Excluded · {row.get('Reason', '')}", int(row.get("Rows", 0) or 0)))

    return pd.DataFrame(rows, columns=["Field", "Value"])


# ---------- Sentiment helpers ----------

def add_final_sentiment_columns(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()

    out = df.copy()

    base_ai_sent = out.get("AI Sentiment", pd.Series(index=out.index, dtype="object"))
    base_ai_conf = out.get("AI Sentiment Confidence", pd.Series(index=out.index, dtype="object"))
    base_ai_rat = out.get("AI Sentiment Rationale", pd.Series(index=out.index, dtype="object"))

    review_ai_sent = out.get("Review AI Sentiment", pd.Series(index=out.index, dtype="object"))
    review_ai_conf = out.get("Review AI Confidence", pd.Series(index=out.index, dtype="object"))
    review_ai_rat = out.get("Review AI Rationale", pd.Series(index=out.index, dtype="object"))

    review_sent_clean = review_ai_sent.fillna("").astype(str).str.strip()

    out["AI Sentiment"] = review_ai_sent.where(review_sent_clean != "", base_ai_sent)
    out["AI Sentiment Confidence"] = review_ai_conf.where(review_sent_clean != "", base_ai_conf)
    out["AI Sentiment Rationale"] = review_ai_rat.where(review_sent_clean != "", base_ai_rat)

    assigned = out.get("Assigned Sentiment", pd.Series(index=out.index, dtype="object"))
    ai = out.get("AI Sentiment", pd.Series(index=out.index, dtype="object"))

    assigned_clean = assigned.fillna("").astype(str).str.strip()
    ai_clean = ai.fillna("").astype(str).str.strip()

    out["Hybrid Sentiment"] = assigned_clean.where(assigned_clean != "", ai_clean)
    out["Hybrid Sentiment"] = out["Hybrid Sentiment"].replace("", pd.NA)

    if "AI Sentiment Confidence" in out.columns:
        out["Hybrid Sentiment Confidence"] = out["AI Sentiment Confidence"]
        out.loc[assigned_clean != "", "Hybrid Sentiment Confidence"] = pd.NA

    return out


def _docx_add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _docx_add_example_block(
    document,
    headline: str,
    url: str = "",
    metrics_line: str = "",
) -> None:
    headline = str(headline or "").strip()
    metrics_line = str(metrics_line or "").strip()
    url = str(url or "").strip()
    if not headline:
        return

    p = document.add_paragraph()
    p.paragraph_format.space_after = 0
    if url:
        _docx_add_hyperlink(p, headline, url)
    else:
        p.add_run(headline)

    if metrics_line:
        m = document.add_paragraph(metrics_line)
        m.paragraph_format.space_before = 0
        m.paragraph_format.space_after = 0


def _format_metric_parts(parts: list[tuple[str, Any]]) -> str:
    formatted: list[str] = []
    for label, value in parts:
        if value is None or value == "":
            continue
        if isinstance(value, float):
            if pd.isna(value):
                continue
            if value.is_integer():
                value = int(value)
        formatted.append(f"{label}: {value:,}" if isinstance(value, int) else f"{label}: {value}")
    return " | ".join(formatted)


def _safe_string(value: Any) -> str:
    return str(value or "").strip()


def _iter_author_report_blocks(session_state) -> list[dict[str, Any]]:
    df_traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
    auth_outlet_table = session_state.get("auth_outlet_table", pd.DataFrame()).copy()
    auth_outlet_table = auth_outlet_table if isinstance(auth_outlet_table, pd.DataFrame) and not auth_outlet_table.empty else None
    summary_df, story_level_df = build_author_metrics(df_traditional, auth_outlet_table=auth_outlet_table)
    if summary_df.empty:
        return []

    selected_authors = [
        _safe_string(author)
        for author in session_state.get("author_insights_selected_authors", [])
        if _safe_string(author)
    ]
    summaries = dict(session_state.get("author_insights_summaries", {}))
    authors = selected_authors or [_safe_string(author) for author in summaries.keys() if _safe_string(author)]
    blocks: list[dict[str, Any]] = []

    for author in authors:
        author_row_df = summary_df[summary_df["Author"] == author]
        if author_row_df.empty:
            continue
        author_row = author_row_df.iloc[0]
        examples_df = build_author_headline_table(story_level_df, author, limit=5)
        blocks.append(
            {
                "title": author,
                "subtitle": _safe_string(author_row.get("Assigned Outlet", "")),
                "summary": _safe_string(summaries.get(author, "")),
                "metrics": _format_metric_parts(
                    [
                        ("Mentions", int(author_row.get("Mention_Total", 0) or 0)),
                        ("Unique Mentions", int(author_row.get("Unique_Stories", 0) or 0)),
                        ("Impressions", int(author_row.get("Impressions", 0) or 0)),
                        ("Effective Reach", int(author_row.get("Effective_Reach", 0) or 0)),
                    ]
                ),
                "examples": [
                    {
                        "headline": _safe_string(row.get("Headline", "")),
                        "url": _safe_string(row.get("Representative URL", "")),
                        "metrics": _format_metric_parts(
                            [
                                ("Outlet", _safe_string(row.get("Representative Outlet", ""))),
                                ("Media Type", _safe_string(row.get("Type", ""))),
                                ("Mentions", int(row.get("Story Mentions", 0) or 0)),
                                ("Impressions", int(row.get("Story Impressions", 0) or 0)),
                                ("Effective Reach", int(row.get("Story Effective Reach", 0) or 0)),
                            ]
                        ),
                    }
                    for _, row in examples_df.iterrows()
                    if _safe_string(row.get("Headline", ""))
                ],
            }
        )

    return blocks


def _iter_outlet_report_blocks(session_state) -> list[dict[str, Any]]:
    df_traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
    summary_df, story_level_df = build_outlet_metrics(
        df_traditional,
        outlet_rollup_map=session_state.get("outlet_rollup_map", {}),
    )
    if summary_df.empty:
        return []

    selected_outlets = [
        _safe_string(outlet)
        for outlet in session_state.get("outlet_insights_selected_outlets", [])
        if _safe_string(outlet)
    ]
    summaries = dict(session_state.get("outlet_insights_summaries", {}))
    outlets = selected_outlets or [_safe_string(outlet) for outlet in summaries.keys() if _safe_string(outlet)]
    blocks: list[dict[str, Any]] = []

    for outlet in outlets:
        outlet_row_df = summary_df[summary_df["Outlet"] == outlet]
        if outlet_row_df.empty:
            continue
        outlet_row = outlet_row_df.iloc[0]
        examples_df = build_outlet_headline_table(story_level_df, outlet, limit=5)
        blocks.append(
            {
                "title": outlet,
                "subtitle": _safe_string(outlet_row.get("Top_Types", "")),
                "summary": _safe_string(summaries.get(outlet, "")),
                "metrics": _format_metric_parts(
                    [
                        ("Mentions", int(outlet_row.get("Mention_Total", 0) or 0)),
                        ("Unique Mentions", int(outlet_row.get("Unique_Mentions", 0) or 0)),
                        ("Impressions", int(outlet_row.get("Impressions", 0) or 0)),
                        ("Effective Reach", int(outlet_row.get("Effective_Reach", 0) or 0)),
                    ]
                ),
                "examples": [
                    {
                        "headline": _safe_string(row.get("Headline", "")),
                        "url": _safe_string(row.get("Representative URL", "")),
                        "metrics": _format_metric_parts(
                            [
                                ("Author", _safe_string(row.get("Author", ""))),
                                ("Media Type", _safe_string(row.get("Type", ""))),
                                ("Mentions", int(row.get("Story Mentions", 0) or 0)),
                                ("Impressions", int(row.get("Story Impressions", 0) or 0)),
                                ("Effective Reach", int(row.get("Story Effective Reach", 0) or 0)),
                            ]
                        ),
                    }
                    for _, row in examples_df.iterrows()
                    if _safe_string(row.get("Headline", ""))
                ],
            }
        )

    return blocks


def _iter_sentiment_report_sections(session_state) -> tuple[str, list[dict[str, Any]]]:
    observation_output = session_state.get("sentiment_observation_output", {}) or {}
    overall = _safe_string(observation_output.get("overall_observation", ""))
    sections = []
    for section in observation_output.get("sentiment_sections", []) or []:
        label = _safe_string(section.get("sentiment", ""))
        if not label:
            continue
        examples = []
        for item in (observation_output.get("_examples_by_sentiment", {}) or {}).get(label, []):
            examples.append(
                {
                    "headline": _safe_string(item.get("headline", "")),
                    "url": _safe_string(item.get("url", "")),
                    "metrics": _format_metric_parts(
                        [
                            ("Outlet", _safe_string(item.get("outlet", ""))),
                            ("Media Type", _safe_string(item.get("example_type", ""))),
                            ("Mentions", int(item.get("mentions", 0) or 0)),
                            ("Impressions", int(item.get("impressions", 0) or 0)),
                            ("Effective Reach", int(item.get("effective_reach", 0) or 0)),
                        ]
                    ),
                }
            )
        sections.append(
            {
                "title": label,
                "summary": _safe_string(section.get("observation", "")),
                "examples": examples,
            }
        )
    return overall, sections


def _iter_regions_report_sections(session_state) -> list[dict[str, Any]]:
    if not bool(session_state.get("regions_prepared", False)):
        return []

    df_traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
    if df_traditional.empty:
        return []

    prepared = build_regions_source_df(
        df_traditional,
        outlet_rollup_map=session_state.get("outlet_rollup_map", {}),
    )
    qualitative_flags = tuple(get_qualitative_coverage_flag_exclusions(session_state))
    prepared = apply_session_coverage_flag_policy(prepared, session_state, list(qualitative_flags))
    filtered_df = filter_regions_df(
        prepared,
        exclude_coverage_flags=[],
        include_countries=session_state.get("regions_include_countries", []),
        exclude_countries=session_state.get("regions_exclude_countries", []),
    )
    current_signature = _build_regions_signature_for_export(session_state, filtered_df)
    generated_store = session_state.get("regions_generated_output", {}) or {}

    selected_levels = session_state.get("regions_analysis_levels", ["Countries", "States / Provinces", "Cities"])
    level_labels = [label for label in ["Countries", "States / Provinces", "Cities"] if label in selected_levels]
    sections: list[dict[str, Any]] = []

    for label in level_labels:
        level_entry = generated_store.get(label, {})
        level_copy: dict[str, Any] = {}
        if isinstance(level_entry, dict) and "content" in level_entry and level_entry.get("signature") == current_signature:
            level_copy = level_entry.get("content", {}) or {}
        elif isinstance(level_entry, dict) and "content" not in level_entry:
            level_copy = level_entry
        if not level_copy:
            continue

        sections.append(
            {
                "title": label,
                "overall": _safe_string(level_copy.get("overall_observation", "")),
                "profiles": [
                    {
                        "region": _safe_string(item.get("region", "")),
                        "blurb": _safe_string(item.get("blurb", "")),
                    }
                    for item in level_copy.get("top_region_profiles", []) or []
                    if _safe_string(item.get("region", "")) and _safe_string(item.get("blurb", ""))
                ],
                "tail": _safe_string(level_copy.get("tail_observation", "")),
            }
        )

    return sections


def _iter_tag_report_sections(session_state) -> tuple[str, list[dict[str, Any]]]:
    observation_output = session_state.get("tagging_observation_output", {}) or {}
    overall = _safe_string(observation_output.get("overall_observation", ""))
    sections = []
    for section in observation_output.get("tag_sections", []) or []:
        label = _safe_string(section.get("tag", ""))
        if not label:
            continue
        examples = []
        for item in (observation_output.get("_examples_by_tag", {}) or {}).get(label, []):
            examples.append(
                {
                    "headline": _safe_string(item.get("headline", "")),
                    "url": _safe_string(item.get("url", "")),
                    "metrics": _format_metric_parts(
                        [
                            ("Outlet", _safe_string(item.get("outlet", ""))),
                            ("Media Type", _safe_string(item.get("example_type", ""))),
                            ("Mentions", int(item.get("mentions", 0) or 0)),
                            ("Impressions", int(item.get("impressions", 0) or 0)),
                            ("Effective Reach", int(item.get("effective_reach", 0) or 0)),
                        ]
                    ),
                }
            )
        sections.append(
            {
                "title": label,
                "summary": _safe_string(section.get("observation", "")),
                "examples": examples,
            }
        )
    return overall, sections


def _iter_top_story_blocks(session_state) -> tuple[str, list[dict[str, Any]]]:
    top_stories = normalize_summary_df(session_state.get("added_df", pd.DataFrame()).copy())
    if top_stories.empty:
        return "", []

    overall = _safe_string((session_state.get("top_story_observation_output", {}) or {}).get("overall_observation", ""))
    top_stories = top_stories.sort_values(["Mentions", "Impressions"], ascending=False)
    blocks = []
    for _, row in top_stories.iterrows():
        headline = _safe_string(row.get("Headline", ""))
        summary = _safe_string(row.get("Top Story Summary", "")) or _safe_string(row.get("Chart Callout", ""))
        if not headline:
            continue
        blocks.append(
            {
                "title": headline,
                "url": _safe_string(row.get("Example URL", "")),
                "summary": summary,
                "metrics": _format_metric_parts(
                    [
                        ("Outlet", _safe_string(row.get("Example Outlet", ""))),
                        ("Media Type", _safe_string(row.get("Example Type", ""))),
                        ("Mentions", int(row.get("Mentions", 0) or 0)),
                        ("Impressions", int(row.get("Impressions", 0) or 0)),
                    ]
                ),
            }
        )
    return overall, blocks


def build_report_copy_docx_bytes(session_state) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to build the report copy document.") from exc

    document = Document()
    document.add_heading("Report Copy", level=0)
    client_name = _safe_string(session_state.get("client_name", "")) or "Client"
    document.add_paragraph(f"{client_name} report copy export")
    document.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    has_content = False

    author_blocks = _iter_author_report_blocks(session_state)
    if author_blocks:
        has_content = True
        document.add_heading("Top Authors", level=1)
        for block in author_blocks:
            title = block["title"]
            subtitle = _safe_string(block.get("subtitle", ""))
            document.add_heading(f"{title} | {subtitle}" if subtitle else title, level=2)
            if block.get("summary"):
                document.add_paragraph(block["summary"])
            if block.get("metrics"):
                document.add_paragraph(block["metrics"])
            if block.get("examples"):
                p = document.add_paragraph()
                p.add_run("Representative examples").bold = True
                for example in block["examples"]:
                    _docx_add_example_block(document, example["headline"], example.get("url", ""), example.get("metrics", ""))

    outlet_blocks = _iter_outlet_report_blocks(session_state)
    if outlet_blocks:
        has_content = True
        document.add_heading("Top Outlets", level=1)
        for block in outlet_blocks:
            title = block["title"]
            subtitle = _safe_string(block.get("subtitle", ""))
            document.add_heading(f"{title} | {subtitle}" if subtitle else title, level=2)
            if block.get("summary"):
                document.add_paragraph(block["summary"])
            if block.get("metrics"):
                document.add_paragraph(block["metrics"])
            if block.get("examples"):
                p = document.add_paragraph()
                p.add_run("Representative examples").bold = True
                for example in block["examples"]:
                    _docx_add_example_block(document, example["headline"], example.get("url", ""), example.get("metrics", ""))

    top_story_overall, top_story_blocks = _iter_top_story_blocks(session_state)
    if top_story_overall or top_story_blocks:
        has_content = True
        document.add_heading("Top Stories", level=1)
        if top_story_overall:
            document.add_heading("Overall Observations", level=2)
            document.add_paragraph(top_story_overall)
        for block in top_story_blocks:
            document.add_heading(block["title"], level=2)
            if block.get("summary"):
                document.add_paragraph(block["summary"])
            if block.get("metrics"):
                document.add_paragraph(block["metrics"])

    sentiment_overall, sentiment_sections = _iter_sentiment_report_sections(session_state)
    if sentiment_overall or sentiment_sections:
        has_content = True
        document.add_heading("Sentiment Insights", level=1)
        if sentiment_overall:
            document.add_heading("Overall Observations", level=2)
            document.add_paragraph(sentiment_overall)
        for section in sentiment_sections:
            document.add_heading(section["title"], level=2)
            if section.get("summary"):
                document.add_paragraph(section["summary"])
            if section.get("examples"):
                p = document.add_paragraph()
                p.add_run("Representative examples").bold = True
                for example in section["examples"]:
                    _docx_add_example_block(document, example["headline"], example.get("url", ""), example.get("metrics", ""))

    tag_overall, tag_sections = _iter_tag_report_sections(session_state)
    if tag_overall or tag_sections:
        has_content = True
        document.add_heading("Tag Insights", level=1)
        if tag_overall:
            document.add_heading("Overall Observations", level=2)
            document.add_paragraph(tag_overall)
        for section in tag_sections:
            document.add_heading(section["title"], level=2)
            if section.get("summary"):
                document.add_paragraph(section["summary"])
            if section.get("examples"):
                p = document.add_paragraph()
                p.add_run("Representative examples").bold = True
                for example in section["examples"]:
                    _docx_add_example_block(document, example["headline"], example.get("url", ""), example.get("metrics", ""))

    regions_sections = _iter_regions_report_sections(session_state)
    if regions_sections:
        has_content = True
        document.add_heading("Regions", level=1)
        for section in regions_sections:
            document.add_heading(section["title"], level=2)
            if section.get("overall"):
                document.add_heading("Overall Observation", level=3)
                document.add_paragraph(section["overall"])
            if section.get("profiles"):
                document.add_heading("Top Region Profiles", level=3)
                for item in section["profiles"]:
                    document.add_paragraph(f"{item['region']}: {item['blurb']}")
            if section.get("tail"):
                document.add_heading("Tail Observation", level=3)
                document.add_paragraph(section["tail"])

    if not has_content:
        raise ValueError("No report-copy content is available yet. Generate at least one set of insights or observations first.")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# ---------- Workbook builder ----------

def build_clean_workbook_bytes(session_state) -> bytes:
    traditional, excluded_rows, excluded_counts = build_scoped_traditional_export_bundle(session_state)
    social = session_state.get("df_social", pd.DataFrame()).copy()
    top_stories = session_state.get("added_df", pd.DataFrame()).copy()
    dupes = session_state.get("df_dupes", pd.DataFrame()).copy()
    raw = session_state.get("df_untouched", pd.DataFrame()).copy()
    outlet_rollup_map = session_state.get("outlet_rollup_map", {})

    for df in [traditional, social, dupes]:
        if isinstance(df, pd.DataFrame) and "Published Date" in df.columns:
            df.drop(columns=["Published Date"], inplace=True)

    original_ave_col = session_state.get("original_ave_col") or "AVE"

    traditional = merge_full_scope_ai_columns_into_clean_trad(session_state, traditional)
    traditional = add_final_sentiment_columns(traditional)
    traditional = explode_tags(traditional)

    social = explode_tags(social)

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter", datetime_format="yyyy-mm-dd") as writer:
        workbook = writer.book
        cleaned_exports: list[tuple[str, pd.DataFrame, Any]] = []

        number_format = workbook.add_format({"num_format": "#,##0"})
        currency_symbol = get_currency_symbol(original_ave_col)
        currency_format = workbook.add_format({"num_format": f"{currency_symbol}#,##0.00"})

        # CLEAN TRAD
        if len(traditional) > 0:
            trad_export = rename_ave(traditional.copy(), original_ave_col=original_ave_col)
            trad_export = add_mapped_outlet_column(trad_export, outlet_rollup_map=outlet_rollup_map)
            if "Impressions" in trad_export.columns:
                trad_export = trad_export.sort_values(by=["Impressions"], ascending=False)
            trad_export.to_excel(writer, sheet_name="CLEAN TRAD", startrow=1, header=False, index=False)
            ws = writer.sheets["CLEAN TRAD"]
            ws.set_tab_color("black")
            cleaned_exports.append(("CLEAN TRAD", trad_export, ws))

        # CLEAN SOCIAL
        if len(social) > 0:
            social_export = rename_ave(social.copy(), original_ave_col=original_ave_col)
            social_export = add_mapped_outlet_column(social_export, outlet_rollup_map=outlet_rollup_map)
            if "Impressions" in social_export.columns:
                social_export = social_export.sort_values(by=["Impressions"], ascending=False)
            social_export.to_excel(writer, sheet_name="CLEAN SOCIAL", startrow=1, header=False, index=False)
            ws = writer.sheets["CLEAN SOCIAL"]
            ws.set_tab_color("black")
            cleaned_exports.append(("CLEAN SOCIAL", social_export, ws))

        # TAGGING SAMPLE
        if (
            tagging_exists(session_state)
            and not should_merge_tagging_into_clean_trad(session_state)
            and tagging_processed_complete_enough(session_state)
        ):
            tagging_export = rename_ave(build_tagging_sample_export(session_state), original_ave_col=original_ave_col)
            if not tagging_export.empty:
                tagging_export.to_excel(writer, sheet_name="TAGGING SAMPLE", header=True, index=False)
                ws = writer.sheets["TAGGING SAMPLE"]
                ws.set_tab_color("#7f8c8d")
                cleaned_exports.append(("TAGGING SAMPLE", tagging_export, ws))

        # SENTIMENT SAMPLE
        if (
            sentiment_exists(session_state)
            and not should_merge_sentiment_into_clean_trad(session_state)
            and sentiment_processed_complete_enough(session_state)
        ):
            sentiment_export = add_final_sentiment_columns(
                rename_ave(build_sentiment_sample_export(session_state), original_ave_col=original_ave_col)
            )
            if not sentiment_export.empty:
                sentiment_export.to_excel(writer, sheet_name="SENTIMENT SAMPLE", header=True, index=False)
                ws = writer.sheets["SENTIMENT SAMPLE"]
                ws.set_tab_color("#7f8c8d")
                cleaned_exports.append(("SENTIMENT SAMPLE", sentiment_export, ws))

        # AUTHORS
        authors = build_author_insights_export_table(session_state, df_traditional=traditional)
        if authors.empty:
            authors = build_authors_export_table(
                traditional.copy(),
                existing_assignments=session_state.get("auth_outlet_table", pd.DataFrame()).copy()
                if len(session_state.get("auth_outlet_table", pd.DataFrame())) > 0 else None,
            )
        if len(authors) > 0:
            selected_authors = [
                str(author).strip()
                for author in session_state.get("author_insights_selected_authors", [])
                if str(author).strip()
            ]
            if not selected_authors:
                authors = authors.sort_values(by=["Mentions", "Impressions"], ascending=False).copy()
            authors.to_excel(writer, sheet_name="Authors", header=True, index=False)
            ws = writer.sheets["Authors"]
            ws.set_tab_color("green")
            cleaned_exports.append(("Authors", authors, ws))

        # OUTLETS
        outlets = build_outlets_export_table(session_state, df_traditional=traditional)
        if len(outlets) > 0:
            outlets.to_excel(writer, sheet_name="Outlets", header=True, index=False)
            ws = writer.sheets["Outlets"]
            ws.set_tab_color("green")
            cleaned_exports.append(("Outlets", outlets, ws))

        # REGIONS
        regions_export = build_regions_export_table(session_state)
        if len(regions_export) > 0:
            regions_export.to_excel(writer, sheet_name="Regions", header=True, index=False)
            ws = writer.sheets["Regions"]
            ws.set_tab_color("green")
            cleaned_exports.append(("Regions", regions_export, ws))

        # TOP STORIES
        if len(top_stories) > 0:
            top_stories_export = top_stories.copy()
            if all(c in top_stories_export.columns for c in ["Mentions", "Impressions"]):
                top_stories_export = top_stories_export.sort_values(by=["Mentions", "Impressions"], ascending=False)

            desired_top_story_columns = [
                "Headline",
                "Date",
                "Mentions",
                "Impressions",
                "Example Outlet",
                "Example URL",
                "Chart Callout",
                "Top Story Summary",
                "Entity Sentiment",
            ]
            existing_top_story_columns = [c for c in desired_top_story_columns if c in top_stories_export.columns]
            top_stories_export = top_stories_export[existing_top_story_columns].copy()

            top_stories_export.to_excel(writer, sheet_name="Top Stories", header=True, index=False)
            ws = writer.sheets["Top Stories"]
            ws.set_tab_color("green")
            cleaned_exports.append(("Top Stories", top_stories_export, ws))

        # DLTD DUPES
        if len(dupes) > 0:
            dupes_export = rename_ave(dupes.copy(), original_ave_col=original_ave_col)
            dupes_export = add_mapped_outlet_column(dupes_export, outlet_rollup_map=outlet_rollup_map)
            dupes_export.to_excel(writer, sheet_name="DLTD DUPES", header=True, index=False)
            ws = writer.sheets["DLTD DUPES"]
            ws.set_tab_color("#c26f4f")
            cleaned_exports.append(("DLTD DUPES", dupes_export, ws))

        # EXCLUDED ROWS
        if len(excluded_rows) > 0:
            excluded_export = rename_ave(excluded_rows.copy(), original_ave_col=original_ave_col)
            excluded_export = add_mapped_outlet_column(excluded_export, outlet_rollup_map=outlet_rollup_map)
            if "Impressions" in excluded_export.columns:
                excluded_export = excluded_export.sort_values(by=["Impressions"], ascending=False)
            excluded_export.to_excel(writer, sheet_name="EXCLUDED ROWS", header=True, index=False)
            ws = writer.sheets["EXCLUDED ROWS"]
            ws.set_tab_color("#c26f4f")
            cleaned_exports.append(("EXCLUDED ROWS", excluded_export, ws))

        # RAW
        raw_export = rename_ave(raw.copy(), original_ave_col=original_ave_col)
        raw_export = add_mapped_outlet_column(raw_export, outlet_rollup_map=outlet_rollup_map)
        raw_export.drop(["Mentions"], axis=1, inplace=True, errors="ignore")
        raw_export.to_excel(writer, sheet_name="RAW", header=True, index=False)
        ws = writer.sheets["RAW"]
        ws.set_tab_color("#c26f4f")
        cleaned_exports.append(("RAW", raw_export, ws))

        # EXPORT METADATA
        metadata_export = build_export_metadata_sheet(
            session_state,
            scoped_traditional_rows=len(traditional),
            excluded_rows_df=excluded_rows,
            excluded_counts_df=excluded_counts,
        )
        metadata_export.to_excel(writer, sheet_name="EXPORT METADATA", header=True, index=False)
        ws = writer.sheets["EXPORT METADATA"]
        ws.set_tab_color("#4f81bd")
        cleaned_exports.append(("EXPORT METADATA", metadata_export, ws))

        # Apply Excel table structures + formatting
        row_level_sheets = {"CLEAN TRAD", "CLEAN SOCIAL", "TAGGING SAMPLE", "SENTIMENT SAMPLE", "DLTD DUPES"}

        for sheet_name, clean_df, ws in cleaned_exports:
            if clean_df.empty:
                continue

            max_row, max_col = clean_df.shape
            column_settings = [{"header": column} for column in clean_df.columns]
            ws.add_table(0, 0, max_row, max_col - 1, {"columns": column_settings})

            apply_sheet_column_formats(
                worksheet=ws,
                df=clean_df,
                number_format=number_format,
                currency_format=currency_format,
                original_ave_col=original_ave_col,
            )

            if "Mentions" in clean_df.columns and sheet_name in row_level_sheets:
                col_idx = clean_df.columns.get_loc("Mentions")
                ws.set_column(col_idx, col_idx, None, None, {"hidden": True})

    output.seek(0)
    return output.getvalue()
# from __future__ import annotations
#
# import io
# from datetime import datetime
#
# import pandas as pd
#
#
# # ---------- Core helpers ----------
#
# def rename_ave(df: pd.DataFrame, original_ave_col: str | None = None) -> pd.DataFrame:
#     """Restore internal AVE column to original uploaded AVE column name for export."""
#     export_ave_name = original_ave_col or "AVE"
#
#     if "AVE" in df.columns:
#         return df.rename(columns={"AVE": export_ave_name})
#     return df
#
#
# def explode_tags(df: pd.DataFrame) -> pd.DataFrame:
#     """Explode comma-separated Tags to one-hot columns, ignoring blanks."""
#     if df is None or df.empty or "Tags" not in df.columns:
#         return df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()
#
#     out = df.copy()
#
#     tags = out["Tags"].fillna("").astype(str).str.strip()
#     tags = tags.replace({"nan": "", "None": ""})
#     out["Tags"] = tags
#
#     dummies = out["Tags"].str.get_dummies(sep=",")
#     dummies.columns = [c.strip() for c in dummies.columns]
#
#     # drop blank / garbage columns if they somehow appear
#     dummies = dummies.loc[:, [c for c in dummies.columns if c and c.lower() != "nan"]]
#
#     if not dummies.empty:
#         dummies = dummies.astype("category")
#         out = out.join(dummies, how="left", rsuffix=" (tag)")
#
#     return out
#
#
#
#
# def build_authors_export_table(
#     df: pd.DataFrame,
#     existing_assignments: pd.DataFrame | None = None,
# ) -> pd.DataFrame:
#     """
#     Rebuild author summary from current df_traditional and preserve any assigned outlets.
#     Excludes blank author names.
#     """
#     if df is None or df.empty:
#         return pd.DataFrame(columns=["Author", "Outlet", "Mentions", "Impressions"])
#
#     available_cols = [c for c in ["Author", "Mentions", "Impressions"] if c in df.columns]
#     if "Author" not in available_cols:
#         return pd.DataFrame(columns=["Author", "Outlet", "Mentions", "Impressions"])
#
#     working = df[available_cols].copy()
#
#     if "Mentions" not in working.columns:
#         working["Mentions"] = 1
#     if "Impressions" not in working.columns:
#         working["Impressions"] = 0
#
#     working["Author"] = working["Author"].fillna("").astype(str).str.strip()
#     working = working[working["Author"] != ""].copy()
#
#     rebuilt = (
#         working.groupby("Author", as_index=False)[["Mentions", "Impressions"]]
#         .sum()
#     )
#
#     if existing_assignments is not None and len(existing_assignments) > 0 and "Outlet" in existing_assignments.columns:
#         assignment_map = (
#             existing_assignments[["Author", "Outlet"]]
#             .copy()
#             .fillna("")
#         )
#         assignment_map["Author"] = assignment_map["Author"].fillna("").astype(str).str.strip()
#         assignment_map = assignment_map[assignment_map["Author"] != ""].copy()
#         assignment_map = assignment_map.drop_duplicates(subset=["Author"], keep="last")
#
#         rebuilt = rebuilt.merge(assignment_map, on="Author", how="left")
#         rebuilt["Outlet"] = rebuilt["Outlet"].fillna("")
#     else:
#         rebuilt.insert(loc=1, column="Outlet", value="")
#
#     return rebuilt[["Author", "Outlet", "Mentions", "Impressions"]].copy()
#
#
# def get_currency_symbol(original_ave_col: str | None) -> str:
#     """Infer currency symbol from original uploaded AVE column name. Fallback to $."""
#     original_ave_col = str(original_ave_col or "AVE")
#
#     if "(EUR)" in original_ave_col:
#         return "€"
#     if "(GBP)" in original_ave_col:
#         return "£"
#     if "(JPY)" in original_ave_col:
#         return "¥"
#
#     return "$"
#
#
# def apply_sheet_column_formats(
#     worksheet,
#     df: pd.DataFrame,
#     number_format,
#     currency_format,
#     original_ave_col: str | None = None,
# ):
#     """Apply column widths and formats by column name."""
#     if df is None or df.empty:
#         return
#
#     original_ave_col = str(original_ave_col or "AVE")
#
#     column_rules = {
#         "Group ID": {"width": 10},
#         "Date": {"width": 12},
#         "Published Date": {"width": 12},
#         "Author": {"width": 24},
#         "Outlet": {"width": 28},
#         "Headline": {"width": 40},
#         "Title": {"width": 40},
#         "Type": {"width": 14},
#         "Media Type": {"width": 14},
#         "URL": {"width": 30},
#         "Example URL": {"width": 30},
#         "Snippet": {"width": 55},
#         "Coverage Snippet": {"width": 55},
#         "Summary": {"width": 45},
#         "Content": {"width": 70},
#         "Country": {"width": 14},
#         "Prov/State": {"width": 14},
#         "Language": {"width": 12},
#         "Sentiment": {"width": 12},
#         "Assigned Sentiment": {"width": 16},
#         "AI Sentiment": {"width": 16},
#         "AI Sentiment Confidence": {"width": 12, "format": number_format},
#         "AI Sentiment Rationale": {"width": 45},
#         "AI Tags": {"width": 30},
#         "AI Tag Rationale": {"width": 45},
#         "Tag_Processed": {"width": 12},
#         "Mentions": {"width": 12, "format": number_format},
#         "Impressions": {"width": 14, "format": number_format},
#         "Effective Reach": {"width": 16, "format": number_format},
#         "AVE": {"width": 14, "format": currency_format},
#         original_ave_col: {"width": 14, "format": currency_format},
#         "Example Outlet": {"width": 20},
#         "Chart Callout": {"width": 40},
#         "Top Story Summary": {"width": 55},
#         "Entity Sentiment": {"width": 45},
#         "Field": {"width": 32},
#         "Value": {"width": 40},
#     }
#
#     worksheet.set_default_row(22)
#
#     for col_name in df.columns:
#         col_idx = df.columns.get_loc(col_name)
#         rule = column_rules.get(col_name)
#         if rule:
#             worksheet.set_column(
#                 col_idx,
#                 col_idx,
#                 rule.get("width"),
#                 rule.get("format"),
#             )
#
#     worksheet.freeze_panes(1, 0)
#
#
# # ---------- AI scope / merge logic ----------
#
# def tagging_exists(session_state) -> bool:
#     df = session_state.get("df_tagging_rows", pd.DataFrame())
#     return isinstance(df, pd.DataFrame) and not df.empty
#
#
# def sentiment_exists(session_state) -> bool:
#     df = session_state.get("df_sentiment_rows", pd.DataFrame())
#     return isinstance(df, pd.DataFrame) and not df.empty
#
#
# def tagging_full_scope(session_state) -> bool:
#     return (
#         session_state.get("tagging_config_step", False)
#         and session_state.get("tagging_sample_mode") == "full"
#     )
#
#
# def sentiment_full_scope(session_state) -> bool:
#     return (
#         session_state.get("sentiment_config_step", False)
#         and session_state.get("sentiment_sample_mode") == "full"
#     )
#
#
# def tagging_processed_complete_enough(session_state) -> bool:
#     df_unique = session_state.get("df_tagging_unique", pd.DataFrame())
#     if not isinstance(df_unique, pd.DataFrame) or df_unique.empty:
#         return False
#     if "Tag_Processed" not in df_unique.columns:
#         return False
#     return bool(df_unique["Tag_Processed"].fillna(False).all())
#
#
# def sentiment_processed_complete_enough(session_state) -> bool:
#     df_unique = session_state.get("df_sentiment_unique", pd.DataFrame())
#     if not isinstance(df_unique, pd.DataFrame) or df_unique.empty:
#         return False
#     if "AI Sentiment" not in df_unique.columns:
#         return False
#     return bool(df_unique["AI Sentiment"].notna().all())
#
#
# def should_merge_tagging_into_clean_trad(session_state) -> bool:
#     return tagging_exists(session_state) and tagging_full_scope(session_state) and tagging_processed_complete_enough(session_state)
#
#
# def should_merge_sentiment_into_clean_trad(session_state) -> bool:
#     return sentiment_exists(session_state) and sentiment_full_scope(session_state) and sentiment_processed_complete_enough(session_state)
#
#
# def build_tagging_sample_export(session_state) -> pd.DataFrame:
#     df = session_state.get("df_tagging_rows", pd.DataFrame())
#     if not isinstance(df, pd.DataFrame) or df.empty:
#         return pd.DataFrame()
#
#     out = df.copy()
#
#     # Put the most important AI columns near the front if they exist
#     priority_cols = [
#         "Group ID",
#         "Date",
#         "Headline",
#         "Outlet",
#         "Type",
#         "Mentions",
#         "Impressions",
#         "Effective Reach",
#         "AI Tags",
#         "AI Tag Rationale",
#         "Tag_Processed",
#     ]
#
#     existing_priority = [c for c in priority_cols if c in out.columns]
#     remaining_cols = [c for c in out.columns if c not in existing_priority]
#
#     return out[existing_priority + remaining_cols].copy()
#
#
# def build_sentiment_sample_export(session_state) -> pd.DataFrame:
#     df = session_state.get("df_sentiment_rows", pd.DataFrame())
#     if not isinstance(df, pd.DataFrame) or df.empty:
#         return pd.DataFrame()
#
#     out = df.copy()
#
#     priority_cols = [
#         "Group ID",
#         "Date",
#         "Headline",
#         "Outlet",
#         "Type",
#         "Mentions",
#         "Impressions",
#         "Effective Reach",
#         "Assigned Sentiment",
#         "AI Sentiment",
#         "AI Sentiment Confidence",
#         "AI Sentiment Rationale",
#     ]
#
#     existing_priority = [c for c in priority_cols if c in out.columns]
#     remaining_cols = [c for c in out.columns if c not in existing_priority]
#
#     return out[existing_priority + remaining_cols].copy()
#
# # def build_sentiment_sample_export(session_state) -> pd.DataFrame:
# #     df = session_state.get("df_sentiment_rows", pd.DataFrame())
# #     if not isinstance(df, pd.DataFrame) or df.empty:
# #         return pd.DataFrame()
# #
# #     out = df.copy()
# #     desired = [
# #         "Group ID", "Date", "Headline", "Outlet", "Type", "Language",
# #         "Mentions", "Impressions", "Effective Reach",
# #         "Assigned Sentiment", "AI Sentiment", "AI Sentiment Confidence", "AI Sentiment Rationale",
# #         "URL", "Snippet",
# #     ]
# #     existing = [c for c in desired if c in out.columns]
# #     return out[existing].copy()
#
#
# def merge_full_scope_ai_columns_into_clean_trad(session_state, traditional: pd.DataFrame) -> pd.DataFrame:
#     out = traditional.copy()
#
#     if should_merge_tagging_into_clean_trad(session_state):
#         tag_rows = session_state.get("df_tagging_rows", pd.DataFrame()).copy()
#         tag_cols = [c for c in ["Group ID", "AI Tags", "AI Tag Rationale", "Tag_Processed"] if c in tag_rows.columns]
#         if "Group ID" in tag_cols:
#             tag_map = tag_rows[tag_cols].drop_duplicates(subset=["Group ID"], keep="last")
#             cols_to_drop = [c for c in tag_cols if c != "Group ID" and c in out.columns]
#             out = out.drop(columns=cols_to_drop, errors="ignore")
#             out = out.merge(tag_map, on="Group ID", how="left")
#
#     if should_merge_sentiment_into_clean_trad(session_state):
#         sent_rows = session_state.get("df_sentiment_rows", pd.DataFrame()).copy()
#         sent_cols = [c for c in ["Group ID", "Assigned Sentiment", "AI Sentiment", "AI Sentiment Confidence", "AI Sentiment Rationale"] if c in sent_rows.columns]
#         if "Group ID" in sent_cols:
#             sent_map = sent_rows[sent_cols].drop_duplicates(subset=["Group ID"], keep="last")
#             cols_to_drop = [c for c in sent_cols if c != "Group ID" and c in out.columns]
#             out = out.drop(columns=cols_to_drop, errors="ignore")
#             out = out.merge(sent_map, on="Group ID", how="left")
#
#     return out
#
#
# # ---------- Metadata sheet ----------
#
# def build_export_metadata_sheet(session_state) -> pd.DataFrame:
#     original_rows = len(session_state.get("df_untouched", pd.DataFrame()))
#     clean_trad_rows = len(session_state.get("df_traditional", pd.DataFrame()))
#     clean_social_rows = len(session_state.get("df_social", pd.DataFrame()))
#     dupes_rows = len(session_state.get("df_dupes", pd.DataFrame()))
#     row_check = clean_trad_rows + clean_social_rows + dupes_rows
#
#     tag_rows = session_state.get("df_tagging_rows", pd.DataFrame())
#     tag_unique = session_state.get("df_tagging_unique", pd.DataFrame())
#
#     sent_rows = session_state.get("df_sentiment_rows", pd.DataFrame())
#     sent_unique = session_state.get("df_sentiment_unique", pd.DataFrame())
#
#     tag_processed_groups = 0
#     if isinstance(tag_unique, pd.DataFrame) and not tag_unique.empty and "Tag_Processed" in tag_unique.columns:
#         tag_processed_groups = int(tag_unique["Tag_Processed"].fillna(False).sum())
#
#     sent_processed_groups = 0
#     if isinstance(sent_unique, pd.DataFrame) and not sent_unique.empty and "AI Sentiment" in sent_unique.columns:
#         sent_processed_groups = int(sent_unique["AI Sentiment"].notna().sum())
#
#     rows = [
#         ("Export Timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
#         ("Export Name", session_state.get("export_name", "")),
#         ("Client Name", session_state.get("client_name", "")),
#         ("Original Rows", original_rows),
#         ("Clean Traditional Rows", clean_trad_rows),
#         ("Clean Social Rows", clean_social_rows),
#         ("Deleted Duplicate Rows", dupes_rows),
#         ("Row Check Total", row_check),
#         ("Row Check Matches Original", "Yes" if row_check == original_rows else "No"),
#         ("Canonical Group ID Present", "Yes" if "Group ID" in session_state.get("df_traditional", pd.DataFrame()).columns else "No"),
#
#         ("Tagging Run", "Yes" if tagging_exists(session_state) else "No"),
#         ("Tagging Scope", session_state.get("tagging_sample_mode", "")),
#         ("Tagging Sample Size Used", session_state.get("tagging_sample_size", "")),
#         ("Tagging Processed Groups", tag_processed_groups),
#         ("Tagging Rows in Working Set", len(tag_rows) if isinstance(tag_rows, pd.DataFrame) else 0),
#         ("Tagging Merged Into Clean Trad", "Yes" if should_merge_tagging_into_clean_trad(session_state) else "No"),
#
#         ("Sentiment Run", "Yes" if sentiment_exists(session_state) else "No"),
#         ("Sentiment Scope", session_state.get("sentiment_sample_mode", "")),
#         ("Sentiment Sample Size Used", session_state.get("sentiment_sample_size", "")),
#         ("Sentiment Processed Groups", sent_processed_groups),
#         ("Sentiment Rows in Working Set", len(sent_rows) if isinstance(sent_rows, pd.DataFrame) else 0),
#         ("Sentiment Merged Into Clean Trad", "Yes" if should_merge_sentiment_into_clean_trad(session_state) else "No"),
#     ]
#
#     return pd.DataFrame(rows, columns=["Field", "Value"])
#
#
# # ---------- Workbook builder ----------
#
# def build_clean_workbook_bytes(session_state) -> bytes:
#     traditional = session_state.get("df_traditional", pd.DataFrame()).copy()
#     social = session_state.get("df_social", pd.DataFrame()).copy()
#     top_stories = session_state.get("added_df", pd.DataFrame()).copy()
#     dupes = session_state.get("df_dupes", pd.DataFrame()).copy()
#     raw = session_state.get("df_untouched", pd.DataFrame()).copy()
#
#     # Remove Published Date (redundant once Date is normalized)
#     for df in [traditional, social, dupes]:
#         if isinstance(df, pd.DataFrame) and "Published Date" in df.columns:
#             df.drop(columns=["Published Date"], inplace=True)
#
#     original_ave_col = session_state.get("original_ave_col") or "AVE"
#
#     traditional = merge_full_scope_ai_columns_into_clean_trad(session_state, traditional)
#     traditional = add_final_sentiment_columns(traditional)
#     traditional = explode_tags(traditional)
#
#     sentiment_export = add_final_sentiment_columns(
#         rename_ave(build_sentiment_sample_export(session_state), original_ave_col=original_ave_col)
#     )
#
#     social = explode_tags(social)
#
#     output = io.BytesIO()
#
#     with pd.ExcelWriter(output, engine="xlsxwriter", datetime_format="yyyy-mm-dd") as writer:
#         workbook = writer.book
#         cleaned_exports: list[tuple[str, pd.DataFrame, Any]] = []
#
#         number_format = workbook.add_format({"num_format": "#,##0"})
#         currency_symbol = get_currency_symbol(original_ave_col)
#         currency_format = workbook.add_format({"num_format": f"{currency_symbol}#,##0.00"})
#
#         # CLEAN TRAD
#         if len(traditional) > 0:
#             trad_export = rename_ave(traditional.copy(), original_ave_col=original_ave_col)
#             if "Impressions" in trad_export.columns:
#                 trad_export = trad_export.sort_values(by=["Impressions"], ascending=False)
#             trad_export.to_excel(writer, sheet_name="CLEAN TRAD", startrow=1, header=False, index=False)
#             ws = writer.sheets["CLEAN TRAD"]
#             ws.set_tab_color("black")
#             cleaned_exports.append(("CLEAN TRAD", trad_export, ws))
#
#         # CLEAN SOCIAL
#         if len(social) > 0:
#             social_export = rename_ave(social.copy(), original_ave_col=original_ave_col)
#             if "Impressions" in social_export.columns:
#                 social_export = social_export.sort_values(by=["Impressions"], ascending=False)
#             social_export.to_excel(writer, sheet_name="CLEAN SOCIAL", startrow=1, header=False, index=False)
#             ws = writer.sheets["CLEAN SOCIAL"]
#             ws.set_tab_color("black")
#             cleaned_exports.append(("CLEAN SOCIAL", social_export, ws))
#
#         # TAGGING SAMPLE
#         if (
#                 tagging_exists(session_state)
#                 and not should_merge_tagging_into_clean_trad(session_state)
#                 and tagging_processed_complete_enough(session_state)
#         ):
#         # if tagging_exists(session_state) and not should_merge_tagging_into_clean_trad(session_state):
#             tagging_export = rename_ave(build_tagging_sample_export(session_state), original_ave_col=original_ave_col)
#             if not tagging_export.empty:
#                 tagging_export.to_excel(writer, sheet_name="TAGGING SAMPLE", header=True, index=False)
#                 ws = writer.sheets["TAGGING SAMPLE"]
#                 ws.set_tab_color("#7f8c8d")
#                 cleaned_exports.append(("TAGGING SAMPLE", tagging_export, ws))
#
#         # SENTIMENT SAMPLE
#         if (
#                 sentiment_exists(session_state)
#                 and not should_merge_sentiment_into_clean_trad(session_state)
#                 and sentiment_processed_complete_enough(session_state)
#         ):
#         # if sentiment_exists(session_state) and not should_merge_sentiment_into_clean_trad(session_state):
#             sentiment_export = rename_ave(build_sentiment_sample_export(session_state), original_ave_col=original_ave_col)
#             if not sentiment_export.empty:
#                 sentiment_export.to_excel(writer, sheet_name="SENTIMENT SAMPLE", header=True, index=False)
#                 ws = writer.sheets["SENTIMENT SAMPLE"]
#                 ws.set_tab_color("#7f8c8d")
#                 cleaned_exports.append(("SENTIMENT SAMPLE", sentiment_export, ws))
#
#         # AUTHORS
#         authors = build_authors_export_table(
#             session_state.get("df_traditional", pd.DataFrame()).copy(),
#             existing_assignments=session_state.get("auth_outlet_table", pd.DataFrame()).copy()
#             if len(session_state.get("auth_outlet_table", pd.DataFrame())) > 0 else None,
#         )
#         if len(authors) > 0:
#             authors = authors.sort_values(by=["Mentions", "Impressions"], ascending=False).copy()
#             authors.to_excel(writer, sheet_name="Authors", header=True, index=False)
#             ws = writer.sheets["Authors"]
#             ws.set_tab_color("green")
#             cleaned_exports.append(("Authors", authors, ws))
#
#         # TOP STORIES
#         if len(top_stories) > 0:
#             top_stories_export = top_stories.copy()
#             if all(c in top_stories_export.columns for c in ["Mentions", "Impressions"]):
#                 top_stories_export = top_stories_export.sort_values(by=["Mentions", "Impressions"], ascending=False)
#
#             desired_top_story_columns = [
#                 "Headline",
#                 "Date",
#                 "Mentions",
#                 "Impressions",
#                 "Example Outlet",
#                 "Example URL",
#                 "Chart Callout",
#                 "Top Story Summary",
#                 "Entity Sentiment",
#             ]
#             existing_top_story_columns = [c for c in desired_top_story_columns if c in top_stories_export.columns]
#             top_stories_export = top_stories_export[existing_top_story_columns].copy()
#
#             top_stories_export.to_excel(writer, sheet_name="Top Stories", header=True, index=False)
#             ws = writer.sheets["Top Stories"]
#             ws.set_tab_color("green")
#             cleaned_exports.append(("Top Stories", top_stories_export, ws))
#
#         # DLTD DUPES
#         if len(dupes) > 0:
#             dupes_export = rename_ave(dupes.copy(), original_ave_col=original_ave_col)
#             dupes_export.to_excel(writer, sheet_name="DLTD DUPES", header=True, index=False)
#             ws = writer.sheets["DLTD DUPES"]
#             ws.set_tab_color("#c26f4f")
#             cleaned_exports.append(("DLTD DUPES", dupes_export, ws))
#
#         # RAW
#         raw_export = rename_ave(raw.copy(), original_ave_col=original_ave_col)
#         raw_export.drop(["Mentions"], axis=1, inplace=True, errors="ignore")
#         raw_export.to_excel(writer, sheet_name="RAW", header=True, index=False)
#         ws = writer.sheets["RAW"]
#         ws.set_tab_color("#c26f4f")
#         cleaned_exports.append(("RAW", raw_export, ws))
#
#         # EXPORT METADATA
#         metadata_export = build_export_metadata_sheet(session_state)
#         metadata_export.to_excel(writer, sheet_name="EXPORT METADATA", header=True, index=False)
#         ws = writer.sheets["EXPORT METADATA"]
#         ws.set_tab_color("#4f81bd")
#         cleaned_exports.append(("EXPORT METADATA", metadata_export, ws))
#
#         # Apply Excel table structures + formatting
#         for _, clean_df, ws in cleaned_exports:
#             if clean_df.empty:
#                 continue
#
#             max_row, max_col = clean_df.shape
#             column_settings = [{"header": column} for column in clean_df.columns]
#             ws.add_table(0, 0, max_row, max_col - 1, {"columns": column_settings})
#
#             if "Mentions" in clean_df.columns:
#                 col_idx = clean_df.columns.get_loc("Mentions")
#                 ws.set_column(col_idx, col_idx, None, None, {"hidden": True})
#
#             apply_sheet_column_formats(
#                 worksheet=ws,
#                 df=clean_df,
#                 number_format=number_format,
#                 currency_format=currency_format,
#                 original_ave_col=original_ave_col,
#             )
#
#     output.seek(0)
#     return output.getvalue()
#
# def add_final_sentiment_columns(df: pd.DataFrame) -> pd.DataFrame:
#     if df is None or df.empty:
#         return df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()
#
#     out = df.copy()
#
#     # --- Resolve AI fields (prefer review if present) ---
#     ai_sent = None
#     ai_conf = None
#     ai_rat = None
#
#     if "Review AI Sentiment" in out.columns:
#         ai_sent = out["Review AI Sentiment"]
#         ai_conf = out.get("Review AI Confidence")
#         ai_rat = out.get("Review AI Rationale")
#     else:
#         ai_sent = out.get("AI Sentiment")
#         ai_conf = out.get("AI Sentiment Confidence")
#         ai_rat = out.get("AI Sentiment Rationale")
#
#     if ai_sent is not None:
#         out["AI Sentiment"] = ai_sent
#     if ai_conf is not None:
#         out["AI Sentiment Confidence"] = ai_conf
#     if ai_rat is not None:
#         out["AI Sentiment Rationale"] = ai_rat
#
#     # --- Hybrid sentiment ---
#     assigned = out.get("Assigned Sentiment", pd.Series(index=out.index, dtype="object"))
#     ai = out.get("AI Sentiment", pd.Series(index=out.index, dtype="object"))
#
#     assigned_clean = assigned.fillna("").astype(str).str.strip()
#     ai_clean = ai.fillna("").astype(str).str.strip()
#
#     out["Hybrid Sentiment"] = assigned_clean.where(assigned_clean != "", ai_clean)
#     out["Hybrid Sentiment"] = out["Hybrid Sentiment"].replace("", pd.NA)
#
#     # Confidence
#     if "AI Sentiment Confidence" in out.columns:
#         out["Hybrid Sentiment Confidence"] = out["AI Sentiment Confidence"]
#         out.loc[assigned_clean != "", "Hybrid Sentiment Confidence"] = pd.NA
#
#     return out

# def add_final_sentiment_columns(df: pd.DataFrame) -> pd.DataFrame:
#     if df is None or df.empty:
#         return df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()
#
#     out = df.copy()
#
#     assigned = out["Assigned Sentiment"] if "Assigned Sentiment" in out.columns else pd.Series(index=out.index, dtype="object")
#     ai = out["AI Sentiment"] if "AI Sentiment" in out.columns else pd.Series(index=out.index, dtype="object")
#
#     assigned_clean = assigned.fillna("").astype(str).str.strip()
#     ai_clean = ai.fillna("").astype(str).str.strip()
#
#     out["Final Sentiment"] = assigned_clean.where(assigned_clean != "", ai_clean)
#     out["Final Sentiment"] = out["Final Sentiment"].replace("", pd.NA)
#
#     if "AI Sentiment Confidence" in out.columns:
#         out["Final Sentiment Confidence"] = out["AI Sentiment Confidence"]
#         if "Assigned Sentiment" in out.columns:
#             out.loc[assigned_clean != "", "Final Sentiment Confidence"] = pd.NA
#
#     return out
